"""Document parser for CV files — extracts text from PDF, DOCX, and plain text."""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


class CvParser:
    async def parse(self, file_path: str, mime_type: str) -> str:
        fmt = SUPPORTED_MIME_TYPES.get(mime_type)
        if fmt is None:
            raise ValueError(f"Unsupported CV mime type: {mime_type}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            if fmt == "pdf":
                return self._parse_pdf(path)
            elif fmt == "docx":
                return self._parse_docx(path)
            else:
                return path.read_text(encoding="utf-8")
        except Exception as e:
            logger.error("CV parse failed for %s: %s", file_path, e)
            raise RuntimeError(f"Failed to parse CV: {e}") from e

    def _parse_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except ImportError:
            logger.warning("pypdf not installed, falling back to pdfplumber")
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                return "\n\n".join(page.extract_text() or "" for page in pdf.pages)

    def _parse_docx(self, path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(para.text for para in doc.paragraphs)


cv_parser = CvParser()


def extract_text_from_bytes(content: bytes, mime_type: str) -> str:
    """Extract text from file bytes based on MIME type."""
    if mime_type == "text/plain":
        return content.decode("utf-8", errors="replace")

    if mime_type == "application/pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n\n".join(pages).strip()
        except ImportError:
            logger.warning("pdfplumber not installed, falling back to basic extraction")
            return content.decode("utf-8", errors="replace")

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs).strip()
        except ImportError:
            logger.warning("python-docx not installed, falling back to basic extraction")
            return content.decode("utf-8", errors="replace")

    return content.decode("utf-8", errors="replace")


def extract_text_from_path(file_path: str, mime_type: str) -> str:
    """Extract text from a file path."""
    with open(file_path, "rb") as f:
        content = f.read()
    return extract_text_from_bytes(content, mime_type)
